import requests
import json
from docx import Document

# SET url_base variable to REST API endpoint for Batch Speech to text Transcription results
url_base = 'https://eastus.api.cognitive.microsoft.com/speechtotext/v3.2/transcriptions/'

# SET transcriptionid as found in response body from batchspeechtotext.py
transcriptionid='aabbccd-aabb-aabb-aabb-aabbccddeeff'
url =url_base+transcriptionid

# SET Ocp-Apim-Subscription-Key to the 32-character Private Key, which is required to call Azure Rest API endpoint. This key can be found in your Azure console.
headers={
    "Ocp-Apim-Subscription-Key":"aabbccddeeffgghhiijjkkllmmnnooppqq",
    "Content-Type":"application/json"
    }
res = requests.get(url,headers=headers)

# convert response body to json formated string 
jsonres = json.dumps(res.json())

# convert json formated string to python object
data = json.loads(jsonres)

# Check if transcribing is still in progress. If in progress, then wait to run this python script again.
if(data['status']=='Running'):
    print('Transcription still '+data['status']+'. Please wait.')

# If transcribing is finished, then proceed to saving the transcription to a word document.
if(data['status']=='Succeeded'):

    # SET urlfile to REST API endpoint to retrieve transcription files
    urlfile = url+'/files'
    res = requests.get(urlfile,headers=headers)

    # convert response body to json formated string 
    resjson = json.dumps(res.json())

    # convert json formated string to python object
    data = json.loads(resjson)

    # loop over response body to retrieve url to transcription text
    output_dict = [x for x in data['values'] if x['kind'] == 'Transcription']

    # SET urltranscr to url to transcription text
    urltranscr = output_dict[0]['links']['contentUrl']
    restranscrjson = requests.get(urltranscr,headers={"Content-Type":"application/json"})

    # convert response body to json formated string 
    restranscrjsonform = json.dumps(restranscrjson.json())

    # convert json formated string to python object
    dataobj = json.loads(restranscrjsonform)

    # set text variable to the transcription text
    text = dataobj['combinedRecognizedPhrases'][0]['display']

    # initialize a Word document
    document = Document()

    #SET heading for Word document. This helps to name and identify the Word document.
    headingInFile='Test Transcription'
    document.add_heading(headingInFile, 0) 

    #CREATE a new paragraph for the transcription text
    p = document.add_paragraph(text)

    # SET name of file of word document
    nameOfFile='Test-Transcription.docx'

    # Save Word document to a directory or folder 
    folderlocation = '/location/of/folder/'
    document.save(folderlocation+nameOfFile)
